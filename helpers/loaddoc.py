from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import (
    PyPDFLoader,
    TextLoader,
    UnstructuredWordDocumentLoader,
    CSVLoader,
    DirectoryLoader
)
import os
import tempfile
from langchain.embeddings import HuggingFaceBgeEmbeddings
from typing import List, Union, Any
from langchain.schema import Document
import pandas as pd
from pathlib import Path
from docx import Document as DocxDocument
from langchain_core.documents import Document as LCDocs


# Flexible loader for multiple file formats
def load_all_docs(file: Union[str, Path, Any,"st.runtime.uploaded_file_manager.UploadedFile"]) -> List[Document]:
    """
    Universal loader for PDF, TXT, DOCX, CSV.
    Accepts:
      Single file path
      Streamlit UploadedFile
      List of paths/UploadedFiles
    Returns: list of LangChain Document objects
    """

    temp_path = None

    # If it's a Streamlit UploadedFile, save it to a temporary file
    if hasattr(file, "read")and hasattr(file, "name"):  
        suffix = os.path.splitext(file.name)[-1].lower()
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
            tmp.write(file.read())
            temp_path = tmp.name  
        file = temp_path

    file_path = Path(file)
    ext = file_path.suffix.lower()

    if ext == ".pdf":
       docs = PyPDFLoader(str(file_path)).load()

    elif ext == ".txt":
        docs = TextLoader(str(file_path), encoding="utf-8").load()

    elif ext in [".docx", "docs"]:
        doc_file = DocxDocument(str(file_path))
        text = "\n".join([p.text for p in doc_file.paragraphs if p.text.strip()])
        docs = [Document(page_content=text, metadata={"source": str(file_path)})]

    elif ext == ".csv":
        docs = CSVLoader(str(file_path)).load()

    else:
        raise ValueError(f"Unsupported file format: {ext}")

    # Cleanup temp file if created
    if temp_path and os.path.exists(temp_path):
        os.remove(temp_path)

    return docs

# Keep only minimal metadata (source + content)
def filter_to_minimal_docs(docs: List[Document]) -> List[Document]:
    """
    Given a list of Document objects, return a new list of Document objects
    containing only 'source' in metadata and the original page_content.
    """
    minimal_docs: List[Document] = []
    for doc in docs:
        src = doc.metadata.get("source", "")
        minimal_docs.append(
            LCDocs(
                page_content=doc.page_content,
                metadata={"source": src}
            )
        )
    return minimal_docs


# Split the data into text chunks
def text_split(doc: List[LCDocs]) -> List[LCDocs]:
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=500,
        chunk_overlap=50,
        separators=["\n\n", "\n", " ", ""] 
    )
    #text_chunks = text_splitter.split_documents(extracted_data)
    return text_splitter.split_documents(docs)